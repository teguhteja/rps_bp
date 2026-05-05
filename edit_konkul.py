#!/usr/bin/env python3
"""
Generate Kontrak Kuliah DOCX from JSON data.
Usage: python generate_kontrak_kuliah.py -i data.json [-o output.docx] [--template template.docx]
"""

import argparse
import json
import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
import io

def parse_placeholder(placeholder):
    """
    Parse placeholder seperti '{cpmk[0].kode}' menjadi path list.
    Contoh: 'cpmk[0].kode' -> ['cpmk', 0, 'kode']
    """
    if placeholder.startswith('{') and placeholder.endswith('}'):
        placeholder = placeholder[1:-1]
    placeholder = placeholder.replace(' ', '')
    parts = re.split(r'\.|\[|\]', placeholder)
    parts = [p for p in parts if p]
    for i, p in enumerate(parts):
        if p.isdigit():
            parts[i] = int(p)
    return parts

def get_value(data, path_parts):
    """Ambil nilai dari dictionary/list berdasarkan path parts."""
    # Cek dulu apakah path_parts[0] ada langsung di data
    value = data
    try:
        for part in path_parts:
            if isinstance(value, dict):
                value = value.get(part)
            elif isinstance(value, list):
                value = value[part]
            else:
                return None
            if value is None:
                return None
        return value
    except (IndexError, KeyError, TypeError):
        return None

def set_image_in_front_of_text(picture):
    """Mengubah format gambar menjadi 'In Front of Text'."""
    drawing = picture._inline.getparent()
    inline = drawing.find('.//wp:inline', namespaces=drawing.nsmap)
    if inline is None:
        return
    extent = inline.find('.//wp:extent', namespaces=inline.nsmap)
    effectExtent = inline.find('.//wp:effectExtent', namespaces=inline.nsmap)
    docPr = inline.find('.//wp:docPr', namespaces=inline.nsmap)
    cNvGraphicFramePr = inline.find('.//wp:cNvGraphicFramePr', namespaces=inline.nsmap)
    graphic = inline.find('.//a:graphic', namespaces=inline.nsmap)
    if None in [extent, docPr, graphic]:
        return
    anchor_xml = '''
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240"
                   behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
                <wp:posOffset>0</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
                <wp:posOffset>0</wp:posOffset>
            </wp:positionV>
        </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    anchor.append(extent)
    if effectExtent is not None:
        anchor.append(effectExtent)
    wrap_none = parse_xml('<wp:wrapNone xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"/>')
    anchor.append(wrap_none)
    anchor.append(docPr)
    if cNvGraphicFramePr is not None:
        anchor.append(cNvGraphicFramePr)
    anchor.append(graphic)
    drawing.replace(inline, anchor)

def replace_placeholders_in_text(text, data):
    """Ganti semua placeholder dalam string teks."""
    if not text:
        return text
    pattern = r'\{[^{}]+\}'
    def replacer(match):
        placeholder = match.group(0)
        path = parse_placeholder(placeholder)
        value = get_value(data, path)
        if value is None:
            # Coba alternatif: jika ada 'meta' di data (format RPS)
            if 'meta' in data and path and path[0] in data['meta']:
                value = data['meta'][path[0]]
            else:
                return placeholder
        if isinstance(value, list):
            return '\n'.join(str(v) for v in value)
        return str(value)
    return re.sub(pattern, replacer, text)

def replace_placeholders_in_paragraph(paragraph, data):
    """Ganti placeholder dalam satu paragraf, termasuk gambar tanda tangan."""
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return

    # Deteksi placeholder gambar tanda tangan
    sign_match = None
    sign_key = None
    is_small = False
    
    if '{dosen_sign}' in full_text:
        sign_match = '{dosen_sign}'
        sign_key = 'dosen_sign'
    elif '{dosen_sign_small}' in full_text:
        sign_match = '{dosen_sign_small}'
        sign_key = 'dosen_sign'
        is_small = True
    elif '{mahasiswa_sign}' in full_text:
        sign_match = '{mahasiswa_sign}'
        sign_key = 'mahasiswa_sign'
    elif '{mahasiswa_sign_small}' in full_text:
        sign_match = '{mahasiswa_sign_small}'
        sign_key = 'mahasiswa_sign'
        is_small = True

    if sign_match:
        # Ambil path dari data (coba di meta.sign_key atau langsung sign_key)
        image_path = None
        if sign_key in data:
            image_path = data[sign_key]
        elif 'meta' in data and sign_key in data['meta']:
            image_path = data['meta'][sign_key]
        
        img_width = Inches(0.5) if is_small else Inches(1.0)
        if image_path and os.path.exists(image_path):
            for run in paragraph.runs:
                run.text = ''
            run = paragraph.add_run()
            try:
                with Image.open(image_path) as img:
                    if img.mode not in ('RGB', 'RGBA'):
                        img = img.convert('RGBA')
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='PNG')
                    img_byte_arr.seek(0)
                    pic = run.add_picture(img_byte_arr, width=img_width)
                    set_image_in_front_of_text(pic)
            except Exception as e:
                print(f"Warning: Gagal menambahkan gambar {image_path}: {e}")
        else:
            # Kosongkan saja
            for run in paragraph.runs:
                run.text = ''
        return

    new_text = replace_placeholders_in_text(full_text, data)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)

def replace_placeholders_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)

def process_document(template_path, json_data, output_path):
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, json_data)

    for table in doc.tables:
        replace_placeholders_in_table(table, json_data)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, json_data)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, json_data)
        for table in section.header.tables:
            replace_placeholders_in_table(table, json_data)
        for table in section.footer.tables:
            replace_placeholders_in_table(table, json_data)

    doc.save(output_path)
    print(f"Kontrak Kuliah berhasil disimpan ke: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Generate Kontrak Kuliah DOCX from JSON.')
    parser.add_argument('-i', '--input', required=True, help='Input JSON file (data kontrak kuliah)')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional, default: output/filename.docx)')
    parser.add_argument('--template', default='Kontrak_Kuliah_MK_TSI0000.docx',
                        help='Template DOCX (default: Kontrak_Kuliah_MK_TSI0000.docx)')
    args = parser.parse_args()

    if not args.output:
        base_name = os.path.splitext(os.path.basename(args.input))[0]
        os.makedirs('output', exist_ok=True)
        args.output = os.path.join('output', f"{base_name}_kontrak.docx")

    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error membaca JSON: {e}")
        return

    try:
        process_document(args.template, data, args.output)
    except FileNotFoundError:
        print(f"Error: Template '{args.template}' tidak ditemukan.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    main()