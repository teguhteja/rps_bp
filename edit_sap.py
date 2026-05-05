#!/usr/bin/env python3
"""
Generate seluruh SAP (16 pertemuan) dari satu file JSON SAP.
Output: folder {kode_mk}_{nama_mk}_sap, dengan file {kode_mk}_{nama_mk}_sap_{no:02d}.docx
"""

import argparse
import json
import re
import os
from docx import Document

def parse_placeholder(placeholder):
    if placeholder.startswith('{') and placeholder.endswith('}'):
        placeholder = placeholder[1:-1]
    placeholder = placeholder.replace(' ', '')
    parts = re.split(r'\.|\[|\]', placeholder)
    parts = [p for p in parts if p]
    for i, p in enumerate(parts):
        if p.isdigit():
            parts[i] = int(p)
    return parts

def get_value(data, path_parts):
    value = data
    try:
        for part in path_parts:
            if isinstance(value, dict):
                value = value.get(part)
            elif isinstance(value, list):
                value = value[part]
            else:
                return None
            if value is None:
                return None
        return value
    except (IndexError, KeyError, TypeError):
        return None

def replace_placeholders_in_text(text, data):
    if not text:
        return text
    pattern = r'\{[^{}]+\}'
    def replacer(match):
        placeholder = match.group(0)
        path = parse_placeholder(placeholder)
        value = get_value(data, path)
        if value is None:
            return placeholder
        if isinstance(value, list):
            return '\n'.join(str(v) for v in value)
        return str(value)
    return re.sub(pattern, replacer, text)

def replace_placeholders_in_paragraph(paragraph, data):
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return
    new_text = replace_placeholders_in_text(full_text, data)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)

def replace_placeholders_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)

def process_sap(template_path, sap_data, pertemuan_obj, output_path):
    doc = Document(template_path)

    # Gabungkan data meta dan pertemuan_obj menjadi satu dictionary
    context = {**sap_data['meta'], **pertemuan_obj}
    # Tambahkan juga kegiatan dan lainnya secara langsung
    # Karena kegiatan berbentuk nested, pastikan placeholder bisa diakses
    # Misal {pengajar_pendahuluan} akan diakses via context['kegiatan']['pendahuluan']['pengajar']
    # Untuk memudahkan, kita flatten atau langsung gunakan get_value yang bisa nested
    # Kita akan gunakan get_value dengan data = context + meta, tapi lebih mudah: buat dictionary besar
    full_context = {
        **sap_data['meta'],
        **pertemuan_obj
    }
    
    # Map specific template placeholders to JSON data structure
    if 'no' in pertemuan_obj:
        full_context['no_pertemuan'] = pertemuan_obj['no']
    
    if 'detail_cpmk' in pertemuan_obj:
        full_context['sub_cpmk'] = pertemuan_obj['detail_cpmk']
        
    if 'kegiatan' in pertemuan_obj:
        kegiatan = pertemuan_obj['kegiatan']
        for phase in ['pendahuluan', 'penyajian', 'penutup']:
            if phase in kegiatan:
                for actor in ['pengajar', 'mahasiswa', 'media']:
                    if actor in kegiatan[phase]:
                        key = f"{actor}_{phase}"
                        full_context[key] = kegiatan[phase][actor]

    # activities sudah termasuk di pertemuan_obj['kegiatan'], jadi aman

    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, full_context)
    for table in doc.tables:
        replace_placeholders_in_table(table, full_context)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, full_context)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, full_context)

    doc.save(output_path)
    print(f"Generated: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Generate semua SAP dari JSON.')
    parser.add_argument('-i', '--input', required=True, help='Input JSON SAP file')
    parser.add_argument('-o', '--output', default='.', help='Output directory')
    parser.add_argument('--template', default='SAP_MK_TSI0000.docx', help='Template SAP DOCX')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)

    meta = data['meta']
    kode = meta['kode_mk']
    nama = meta['nama_mk'].replace(' ', '_')
    folder_name = os.path.join(args.output, f"{kode}_{nama}_sap")
    os.makedirs(folder_name, exist_ok=True)

    for pertemuan in data['pertemuan']:
        no = pertemuan['no']
        output_file = os.path.join(folder_name, f"{kode}_{nama}_sap_{no:02d}.docx")
        process_sap(args.template, data, pertemuan, output_file)

if __name__ == '__main__':
    main()