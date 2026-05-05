import re
from docx import Document

def extract_placeholders(doc_path):
    doc = Document(doc_path)
    placeholders = set()
    pattern = r'\{[^{}]+\}'

    def find_in_paragraph(p):
        for match in re.finditer(pattern, p.text):
            placeholders.add(match.group(0))

    def find_in_table(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    find_in_paragraph(p)

    for p in doc.paragraphs:
        find_in_paragraph(p)
    for t in doc.tables:
        find_in_table(t)
    for s in doc.sections:
        for p in s.header.paragraphs:
            find_in_paragraph(p)
        for t in s.header.tables:
            find_in_table(t)
        for p in s.footer.paragraphs:
            find_in_paragraph(p)
        for t in s.footer.tables:
            find_in_table(t)

    for p in sorted(list(placeholders)):
        print(p)

if __name__ == '__main__':
    extract_placeholders(r'TSI3217_Matematika_Diskrit_sap\TSI3217_Matematika_Diskrit_sap_01.docx')
