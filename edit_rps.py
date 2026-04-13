#!/usr/bin/env python3
"""
RPS Generator - Mengisi template DOCX dengan data dari JSON.
Usage: python edit_rps_bp.py -i data.json [-o output.docx] [--template template.docx]
"""

import argparse
import json
import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image
import io

def parse_placeholder(placeholder):
    """
    Parse placeholder seperti '{cpl_prodi[0].kode}' menjadi path list.
    Contoh: 'cpl_prodi[0].kode' -> ['cpl_prodi', 0, 'kode']
    """
    # Hapus kurung kurawal dan spasi
    if placeholder.startswith('{') and placeholder.endswith('}'):
        placeholder = placeholder[1:-1]
    placeholder = placeholder.replace(' ', '')
    # Split dengan regex yang menangani tanda kurung siku
    parts = re.split(r'\.|\[|\]', placeholder)
    # Buang string kosong
    parts = [p for p in parts if p]
    # Konversi angka menjadi int
    for i, p in enumerate(parts):
        if p.isdigit():
            parts[i] = int(p)
    return parts

def get_value(data, path_parts):
    """Ambil nilai dari dictionary/list berdasarkan path parts."""
    # Handle meta variables like {nama_mk} which is actually in data['meta']['nama_mk']
    if len(path_parts) == 1 and isinstance(data, dict) and path_parts[0] not in data and 'meta' in data:
        # Mapping 'pangkat' ke 'pangkat_golongan'
        key = 'pangkat_golongan' if path_parts[0] == 'pangkat' else path_parts[0]
        if key in data['meta']:
            return data['meta'][key]

    value = data
    try:
        for part in path_parts:
            if isinstance(value, dict):
                value = value.get(part)
            elif isinstance(value, list):
                value = value[part]
            else:
                return None
            if value is None:
                return None
        return value
    except (IndexError, KeyError, TypeError):
        return None

def set_image_in_front_of_text(picture):
    """Mengubah format gambar menjadi 'In Front of Text' menggunakan modifikasi XML."""
    # Mendapatkan elemen 'inline' dari gambar yang baru ditambahkan
    drawing = picture._inline.getparent()
    inline = drawing.find('.//wp:inline', namespaces=drawing.nsmap)
    
    if inline is None:
        return
    
    # Mendapatkan properti dan konten dari elemen inline
    extent = inline.find('.//wp:extent', namespaces=inline.nsmap)
    effectExtent = inline.find('.//wp:effectExtent', namespaces=inline.nsmap)
    docPr = inline.find('.//wp:docPr', namespaces=inline.nsmap)
    cNvGraphicFramePr = inline.find('.//wp:cNvGraphicFramePr', namespaces=inline.nsmap)
    graphic = inline.find('.//a:graphic', namespaces=inline.nsmap)
    
    if None in [extent, docPr, graphic]:
        return

    # Membuat elemen 'anchor' (floating) baru untuk 'In Front of Text' (behindDoc="0")
    anchor_xml = f'''
        <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" 
                   xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                   distT="0" distB="0" distL="114300" distR="114300" simplePos="0" relativeHeight="251658240"
                   behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="column">
                <wp:posOffset>0</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="paragraph">
                <wp:posOffset>0</wp:posOffset>
            </wp:positionV>
        </wp:anchor>
    '''
    anchor = parse_xml(anchor_xml)
    
    # Memindahkan elemen-elemen dari inline ke anchor dengan urutan spesifik (WAJIB)
    anchor.append(extent)
    if effectExtent is not None:
        anchor.append(effectExtent)
        
    wrap_none = parse_xml('<wp:wrapNone xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"/>')
    anchor.append(wrap_none)
    
    anchor.append(docPr)
    if cNvGraphicFramePr is not None:
        anchor.append(cNvGraphicFramePr)
    anchor.append(graphic)
    
    # Mengganti elemen inline dengan anchor di dalam elemen drawing
    drawing.replace(inline, anchor)

def replace_placeholders_in_text(text, data):
    """Ganti semua placeholder dalam string teks."""
    if not text:
        return text
    pattern = r'\{[^{}]+\}'
    def replacer(match):
        placeholder = match.group(0)
        path = parse_placeholder(placeholder)
        value = get_value(data, path)
        if value is None:
            # Jika placeholder tidak ditemukan, kembalikan teks aslinya
            return placeholder
        # Jika value adalah list, gabungkan dengan newline
        if isinstance(value, list):
            return '\n'.join(str(v) for v in value)
        return str(value)
    return re.sub(pattern, replacer, text)

def replace_placeholders_in_paragraph(paragraph, data):
    """Ganti placeholder dalam satu paragraf dengan mempertahankan format dasar, atau tambahkan gambar."""
    # Gabungkan semua run menjadi teks utuh
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text.strip():
        return
    
    # Cek khusus untuk image (jika teks mengandung "{dosen_sign}" atau "{dosen_sign_small}")
    is_sign = '{dosen_sign}' in full_text
    is_sign_small = '{dosen_sign_small}' in full_text
    
    if is_sign or is_sign_small:
        image_path = get_value(data, ['dosen_sign'])
        # Tentukan ukuran berdasarkan placeholder yang ditemukan
        img_width = Inches(0.5) if is_sign_small else Inches(1.0)
        
        if image_path and os.path.exists(image_path):
            # Bersihkan teks lama
            for run in paragraph.runs:
                run.text = ''
            # Tambahkan gambar
            run = paragraph.add_run()
            try:
                # docx membutuhkan path absolut atau path yang benar
                abs_path = os.path.abspath(image_path)
                
                # Coba konversi image terlebih dahulu untuk memastikan docx bisa membacanya
                try:
                    with Image.open(abs_path) as img:
                        # Jika gambar memiliki transparansi (RGBA), konversi ke RGB/PNG
                        if img.mode != 'RGB' and img.mode != 'RGBA':
                            img = img.convert('RGBA')
                        img_byte_arr = io.BytesIO()
                        img.save(img_byte_arr, format='PNG')
                        img_byte_arr.seek(0)
                        pic = run.add_picture(img_byte_arr, width=img_width)
                        set_image_in_front_of_text(pic)
                except Exception as img_err:
                    # Fallback
                    pic = run.add_picture(abs_path, width=img_width)  # Sesuaikan ukuran
                    set_image_in_front_of_text(pic)
            except Exception as e:
                # Jika error 'unrecognized image part format', abaikan exception dan jangan hapus teks
                print(f"Warning: Gagal menambahkan gambar {abs_path}: {e}")
            return
        else:
            # Jika tidak ada gambar, kosongkan saja tanda tangannya
            for run in paragraph.runs:
                run.text = ''
            return

    new_text = replace_placeholders_in_text(full_text, data)
    if new_text == full_text:
        return
    
    # Pertahankan run pertama untuk format, hapus run lainnya
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)

def replace_placeholders_in_table(table, data):
    """Ganti placeholder di semua sel tabel."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data)

def process_document(template_path, json_data, output_path):
    """Proses template DOCX dan simpan ke output."""
    doc = Document(template_path)
    
    # Proses paragraf di body
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, json_data)
    
    # Proses tabel
    for table in doc.tables:
        replace_placeholders_in_table(table, json_data)
    
    # Proses header dan footer (jika ada)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_placeholders_in_paragraph(paragraph, json_data)
        for paragraph in section.footer.paragraphs:
            replace_placeholders_in_paragraph(paragraph, json_data)
        for table in section.header.tables:
            replace_placeholders_in_table(table, json_data)
        for table in section.footer.tables:
            replace_placeholders_in_table(table, json_data)
    
    doc.save(output_path)
    print(f"Berhasil menyimpan RPS ke: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Generate RPS DOCX from JSON template.')
    parser.add_argument('-i', '--input', required=True, help='Input JSON file')
    parser.add_argument('-o', '--output', help='Output DOCX file (optional, defaults to output/filename.docx)')
    parser.add_argument('--template', default='RPS_MK_TSI0000.docx', 
                        help='Template DOCX file (default: RPS_MK_TSI0000.docx)')
    args = parser.parse_args()
    
    # Auto-generate output filename if not provided
    if not args.output:
        # Dapatkan nama file tanpa ekstensi dari input (misal: "data.json" -> "data")
        base_name = os.path.splitext(os.path.basename(args.input))[0]
        # Buat folder output jika belum ada
        os.makedirs('output', exist_ok=True)
        args.output = os.path.join('output', f"{base_name}.docx")

    # Baca JSON
    try:
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: File input '{args.input}' tidak ditemukan.")
        return
    except json.JSONDecodeError as e:
        print(f"Error: File JSON tidak valid: {e}")
        return
    
    # Cek template
    try:
        process_document(args.template, data, args.output)
    except FileNotFoundError:
        print(f"Error: File template '{args.template}' tidak ditemukan.")
        return

if __name__ == '__main__':
    main()